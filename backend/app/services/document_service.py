"""
ControlNot v2 - Document Service
Generación de documentos Word con reemplazo de placeholders y formato

Migrado de por_partes.py líneas 1688-1743, 1939-1956
"""
import re
import tempfile
from typing import Dict, List, Tuple
from pathlib import Path
from io import BytesIO
import structlog

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement

logger = structlog.get_logger()


class DocumentGenerator:
    """
    Generador de documentos Word desde templates

    Funcionalidad:
    1. Reemplaza placeholders {{nombre}} con valores
    2. Aplica formato **negrita** automáticamente
    3. Procesa párrafos, tablas, headers, footers
    4. Retorna estadísticas de reemplazo
    """

    # Patrón para formato negrita: **texto**
    BOLD_PATTERN = re.compile(r"\*\*(.*?)\*\*")

    # Valor por defecto para placeholders no encontrados
    DEFAULT_MISSING = "**[NO ENCONTRADO]**"

    def __init__(self):
        """Inicializa el generador de documentos"""
        logger.debug("DocumentGenerator inicializado")

    def _replace_in_paragraph(
        self,
        paragraph,
        responses: Dict[str, str],
        placeholders: List[str]
    ) -> int:
        """
        Reemplaza placeholders en un párrafo

        Args:
            paragraph: Párrafo de python-docx
            responses: Diccionario con valores {placeholder: valor}
            placeholders: Lista de placeholders a buscar

        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0
        text = paragraph.text

        for placeholder in placeholders:
            # Buscar tanto {{placeholder}} como {placeholder}
            patterns = [
                f"{{{{{placeholder}}}}}",  # {{placeholder}}
                f"{{{placeholder}}}"       # {placeholder}
            ]

            for pattern in patterns:
                if pattern in text:
                    # Obtener valor de reemplazo
                    value = responses.get(placeholder, self.DEFAULT_MISSING)

                    # Reemplazar
                    text = text.replace(pattern, str(value))
                    replacements += 1

                    logger.debug(
                        "Placeholder reemplazado en párrafo",
                        placeholder=placeholder,
                        value=value[:50] if len(str(value)) > 50 else value
                    )

        # Actualizar texto del párrafo si hubo cambios
        if replacements > 0:
            paragraph.text = text

        return replacements

    def _replace_in_table(
        self,
        table,
        responses: Dict[str, str],
        placeholders: List[str]
    ) -> int:
        """
        Reemplaza placeholders en una tabla

        Args:
            table: Tabla de python-docx
            responses: Diccionario con valores
            placeholders: Lista de placeholders

        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replacements += self._replace_in_paragraph(
                        paragraph,
                        responses,
                        placeholders
                    )

        return replacements

    def _replace_placeholders_in_document(
        self,
        doc: Document,
        responses: Dict[str, str],
        placeholders: List[str]
    ) -> Dict[str, int]:
        """
        Reemplaza todos los placeholders en el documento

        Args:
            doc: Documento de python-docx
            responses: Valores de reemplazo
            placeholders: Lista de placeholders

        Returns:
            Dict[str, int]: Estadísticas de reemplazo:
                - replaced_in_body: int
                - replaced_in_tables: int
                - replaced_in_headers: int
                - replaced_in_footers: int
                - total_replaced: int
                - missing: int (placeholders no encontrados en responses)
        """
        stats = {
            'replaced_in_body': 0,
            'replaced_in_tables': 0,
            'replaced_in_headers': 0,
            'replaced_in_footers': 0,
            'total_replaced': 0,
            'missing': 0
        }

        # 1. Reemplazar en párrafos del body
        for paragraph in doc.paragraphs:
            stats['replaced_in_body'] += self._replace_in_paragraph(
                paragraph,
                responses,
                placeholders
            )

        # 2. Reemplazar en tablas
        for table in doc.tables:
            stats['replaced_in_tables'] += self._replace_in_table(
                table,
                responses,
                placeholders
            )

        # 3. Reemplazar en headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                stats['replaced_in_headers'] += self._replace_in_paragraph(
                    paragraph,
                    responses,
                    placeholders
                )

            # Headers también pueden tener tablas
            for table in header.tables:
                stats['replaced_in_headers'] += self._replace_in_table(
                    table,
                    responses,
                    placeholders
                )

        # 4. Reemplazar en footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                stats['replaced_in_footers'] += self._replace_in_paragraph(
                    paragraph,
                    responses,
                    placeholders
                )

            # Footers también pueden tener tablas
            for table in footer.tables:
                stats['replaced_in_footers'] += self._replace_in_table(
                    table,
                    responses,
                    placeholders
                )

        # Calcular total
        stats['total_replaced'] = (
            stats['replaced_in_body'] +
            stats['replaced_in_tables'] +
            stats['replaced_in_headers'] +
            stats['replaced_in_footers']
        )

        # Contar placeholders sin valor
        for placeholder in placeholders:
            if placeholder not in responses or responses[placeholder] == self.DEFAULT_MISSING:
                stats['missing'] += 1

        return stats

    def _apply_bold_formatting(self, doc: Document) -> int:
        """
        Aplica formato negrita a texto entre **asteriscos**

        Busca patrones **texto** y los convierte a negrita real

        Args:
            doc: Documento de python-docx

        Returns:
            int: Número de conversiones realizadas
        """
        conversions = 0

        # Procesar párrafos
        for paragraph in doc.paragraphs:
            conversions += self._apply_bold_in_paragraph(paragraph)

        # Procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        conversions += self._apply_bold_in_paragraph(paragraph)

        # Procesar headers y footers
        for section in doc.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                conversions += self._apply_bold_in_paragraph(paragraph)

            # Headers también pueden tener tablas
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            conversions += self._apply_bold_in_paragraph(paragraph)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                conversions += self._apply_bold_in_paragraph(paragraph)

            # Footers también pueden tener tablas
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            conversions += self._apply_bold_in_paragraph(paragraph)

        logger.info("Formato negrita aplicado", conversions=conversions)
        return conversions

    def _apply_bold_in_paragraph(self, paragraph) -> int:
        """
        Aplica negrita en un párrafo específico

        Args:
            paragraph: Párrafo de python-docx

        Returns:
            int: Número de conversiones
        """
        conversions = 0
        text = paragraph.text

        # Buscar patrones **texto**
        matches = list(self.BOLD_PATTERN.finditer(text))

        if not matches:
            return 0

        # Limpiar runs existentes
        for run in paragraph.runs:
            run.text = ""

        # Reconstruir párrafo con formato
        last_end = 0

        for match in matches:
            # Texto antes del match (sin negrita)
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # Texto del match (con negrita)
            bold_text = match.group(1)
            run = paragraph.add_run(bold_text)
            run.bold = True

            last_end = match.end()
            conversions += 1

        # Texto después del último match
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

        return conversions

    def generate_document(
        self,
        template_content: bytes,
        responses: Dict[str, str],
        placeholders: List[str],
        output_filename: str
    ) -> Tuple[bytes, Dict]:
        """
        Genera un documento Word desde un template

        FUNCIÓN PRINCIPAL del servicio

        Args:
            template_content: Contenido del template .docx en bytes
            responses: Diccionario con valores {placeholder: valor}
            placeholders: Lista de placeholders del template
            output_filename: Nombre del archivo de salida

        Returns:
            Tuple[bytes, Dict]:
                - bytes: Contenido del documento generado
                - Dict: Estadísticas de generación

        Example:
            >>> with open('template.docx', 'rb') as f:
            ...     template = f.read()
            >>> responses = {
            ...     'Cliente_Nombre': 'Juan Pérez',
            ...     'Fecha': '15/01/2024'
            ... }
            >>> placeholders = ['Cliente_Nombre', 'Fecha', 'Notario']
            >>> content, stats = generator.generate_document(
            ...     template, responses, placeholders, 'output.docx'
            ... )
            >>> print(stats)
            {'total_replaced': 2, 'missing': 1, ...}
        """
        logger.info(
            "Generando documento",
            output_filename=output_filename,
            placeholders_count=len(placeholders),
            responses_count=len(responses)
        )

        # Crear archivo temporal para el template
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_template:
            temp_template_path = Path(temp_template.name)
            temp_template.write(template_content)
            temp_template.flush()

        try:
            # Abrir template
            doc = Document(temp_template_path)

            # 1. Reemplazar placeholders
            replace_stats = self._replace_placeholders_in_document(
                doc,
                responses,
                placeholders
            )

            # 2. Aplicar formato negrita
            bold_conversions = self._apply_bold_formatting(doc)

            # Guardar en BytesIO
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)

            # Estadísticas completas
            stats = {
                **replace_stats,
                'bold_conversions': bold_conversions,
                'output_filename': output_filename,
                'output_size_bytes': len(output_buffer.getvalue())
            }

            logger.info(
                "Documento generado exitosamente",
                **stats
            )

            return output_buffer.getvalue(), stats

        except Exception as e:
            logger.error(
                "Error al generar documento",
                output_filename=output_filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise

        finally:
            # Limpiar archivo temporal
            try:
                temp_template_path.unlink()
            except Exception as cleanup_error:
                logger.warning(
                    "No se pudo eliminar template temporal",
                    temp_path=str(temp_template_path),
                    error=str(cleanup_error)
                )

    def save_document_to_file(
        self,
        document_content: bytes,
        output_path: Path
    ) -> bool:
        """
        Guarda un documento en el filesystem

        Args:
            document_content: Contenido del documento
            output_path: Ruta donde guardar

        Returns:
            bool: True si se guardó exitosamente
        """
        try:
            # Crear directorio padre si no existe
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Guardar archivo
            with open(output_path, 'wb') as f:
                f.write(document_content)

            logger.info(
                "Documento guardado en disco",
                output_path=str(output_path),
                size_bytes=len(document_content)
            )

            return True

        except Exception as e:
            logger.error(
                "Error al guardar documento",
                output_path=str(output_path),
                error=str(e)
            )
            raise


def generate_document_with_dynamic_placeholders(
    responses: Dict[str, str],
    template_content: bytes,
    placeholders: List[str],
    output_filename: str
) -> Tuple[bytes, Dict]:
    """
    Función pública para generar documentos

    Wrapper sobre DocumentGenerator.generate_document()

    Args:
        responses: Valores de placeholders
        template_content: Template en bytes
        placeholders: Lista de placeholders
        output_filename: Nombre del archivo de salida

    Returns:
        Tuple[bytes, Dict]: Contenido del documento y estadísticas
    """
    generator = DocumentGenerator()
    return generator.generate_document(
        template_content,
        responses,
        placeholders,
        output_filename
    )
