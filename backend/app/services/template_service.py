"""
ControlNot v2 - Template Service
Extrae placeholders de templates Word (.docx)

Migrado de por_partes.py líneas 1458-1502
"""
import re
import time
import tempfile
from typing import List
from pathlib import Path
import os
import structlog
from docx import Document

logger = structlog.get_logger()


class PlaceholderExtractor:
    """
    Extrae placeholders de documentos Word

    Soporta dos formatos:
    - {{placeholder}} - Formato primario
    - {placeholder} - Formato alternativo
    """

    # Patrones de búsqueda para placeholders
    PLACEHOLDER_PATTERNS = [
        r'\{\{([^}]+)\}\}',  # {{placeholder}}
        r'\{([^}]+)\}',      # {placeholder}
    ]

    @staticmethod
    def extract_from_bytes(template_content: bytes) -> List[str]:
        """
        Extrae placeholders de un template en bytes

        Args:
            template_content: Contenido del archivo .docx en bytes

        Returns:
            List[str]: Lista ordenada de placeholders únicos encontrados

        Example:
            >>> content = open('template.docx', 'rb').read()
            >>> placeholders = PlaceholderExtractor.extract_from_bytes(content)
            >>> placeholders
            ['Cliente_Nombre', 'Fecha_Escritura', 'Notario_Numero']
        """
        placeholders = set()
        extract_start = time.time()

        # Crear archivo temporal para procesar
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_path = Path(temp_file.name)
            try:
                # Escribir contenido al archivo temporal
                temp_file.write(template_content)
                temp_file.flush()

                logger.debug(
                    "temp_file_created_for_extraction",
                    temp_path=str(temp_path),
                    template_size_bytes=len(template_content)
                )

                # Abrir documento con python-docx
                doc = Document(temp_path)

                logger.debug(
                    "docx_opened_for_extraction",
                    paragraphs=len(doc.paragraphs),
                    tables=len(doc.tables),
                    sections=len(doc.sections)
                )

                # Extraer texto de diferentes partes del documento
                text_parts = []

                # 1. Párrafos del body
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)

                # 2. Tablas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)

                # 3. Headers (encabezados)
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        text_parts.append(paragraph.text)

                # 4. Footers (pies de página)
                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        text_parts.append(paragraph.text)

                # Buscar placeholders en todo el texto extraído
                all_text = "\n".join(text_parts)

                logger.debug(
                    "text_extraction_complete",
                    text_parts_count=len(text_parts),
                    total_text_length=len(all_text)
                )

                # Contadores por patrón para logging detallado
                pattern_counts = {}
                for i, pattern in enumerate(PlaceholderExtractor.PLACEHOLDER_PATTERNS):
                    pattern_name = "double_curly" if i == 0 else "single_curly"
                    pattern_matches = 0
                    matches = re.finditer(pattern, all_text)
                    for match in matches:
                        placeholder = match.group(1).strip()
                        # Ignorar placeholders vacíos o que empiecen con '{' (duplicados de {{...}})
                        if placeholder and not placeholder.startswith('{'):
                            placeholders.add(placeholder)
                            pattern_matches += 1
                    pattern_counts[pattern_name] = pattern_matches

                extract_duration = (time.time() - extract_start) * 1000
                logger.info(
                    "placeholders_extracted_from_template",
                    total_unique=len(placeholders),
                    double_curly_matches=pattern_counts.get("double_curly", 0),
                    single_curly_matches=pattern_counts.get("single_curly", 0),
                    placeholders_preview=sorted(placeholders)[:10],
                    duration_ms=round(extract_duration, 2)
                )

            except Exception as e:
                logger.error(
                    "Error al extraer placeholders del template",
                    error=str(e),
                    error_type=type(e).__name__
                )
                raise
            finally:
                # Limpiar archivo temporal
                try:
                    temp_path.unlink()
                    logger.debug(
                        "temp_file_cleanup_success",
                        temp_path=str(temp_path)
                    )
                except Exception as cleanup_error:
                    logger.warning(
                        "temp_file_cleanup_failed",
                        temp_path=str(temp_path),
                        error=str(cleanup_error)
                    )

        # Retornar lista ordenada
        return sorted(list(placeholders))

    @staticmethod
    def extract_from_file(file_path: Path) -> List[str]:
        """
        Extrae placeholders de un archivo Word en disco

        Args:
            file_path: Path al archivo .docx

        Returns:
            List[str]: Lista ordenada de placeholders únicos
        """
        start_time = time.time()
        logger.info(
            "extract_from_file_starting",
            file_path=str(file_path)
        )

        with open(file_path, 'rb') as f:
            content = f.read()

        logger.debug(
            "file_content_read",
            file_path=str(file_path),
            size_bytes=len(content)
        )

        result = PlaceholderExtractor.extract_from_bytes(content)

        duration_ms = (time.time() - start_time) * 1000
        logger.info(
            "extract_from_file_complete",
            file_path=str(file_path),
            placeholders_found=len(result),
            duration_ms=round(duration_ms, 2)
        )

        return result


def extract_placeholders_from_template(template_content: bytes) -> List[str]:
    """
    Función pública para extraer placeholders de un template

    Esta es la función principal que debe usarse desde otros módulos.
    Wrapper sobre PlaceholderExtractor.extract_from_bytes()

    Args:
        template_content: Contenido del archivo .docx en bytes

    Returns:
        List[str]: Lista ordenada de placeholders únicos

    Example:
        >>> from app.services.template_service import extract_placeholders_from_template
        >>> with open('template.docx', 'rb') as f:
        ...     content = f.read()
        >>> placeholders = extract_placeholders_from_template(content)
        >>> print(placeholders)
        ['Cliente_Apellido', 'Cliente_Nombre', 'Fecha_Escritura']
    """
    return PlaceholderExtractor.extract_from_bytes(template_content)


def validate_template_content(template_content: bytes) -> bool:
    """
    Valida que el contenido sea un archivo Word válido

    Args:
        template_content: Contenido a validar

    Returns:
        bool: True si es válido, False si no
    """
    if not template_content or len(template_content) == 0:
        logger.warning("Template vacío")
        return False

    # Verificar firma de archivo .docx (ZIP)
    # Los archivos .docx son ZIP con firma PK (50 4B)
    if not template_content.startswith(b'PK'):
        logger.warning("Template no tiene firma de archivo .docx")
        return False

    try:
        # Intentar extraer placeholders (validación completa)
        placeholders = extract_placeholders_from_template(template_content)
        logger.info(
            "Template validado exitosamente",
            placeholders_found=len(placeholders)
        )
        return True
    except Exception as e:
        logger.error(
            "Error al validar template",
            error=str(e),
            error_type=type(e).__name__
        )
        return False
