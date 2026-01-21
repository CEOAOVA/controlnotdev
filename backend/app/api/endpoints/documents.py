"""
ControlNot v2 - Documents Endpoints
Endpoints para categorización y generación de documentos

Rutas:
- GET    /api/documents/categories       - Obtener categorías por tipo
- POST   /api/documents/upload           - Subir documentos categorizados
- POST   /api/documents/generate         - Generar documento final
- GET    /api/documents/download/{id}    - Descargar documento generado
- POST   /api/documents/send-email       - Enviar documento por email

Migración v2:
- Usa SessionManager para almacenamiento en memoria (no dependencias circulares)
- Integra con repositorios para persistencia en PostgreSQL
- Multi-tenant aware con tenant_id
"""
from typing import List, Optional
from uuid import UUID
from datetime import datetime, timezone
import time
import csv
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Depends, Query, Body
from fastapi.responses import StreamingResponse
import structlog
import uuid
from io import BytesIO, StringIO

from app.schemas import (
    CategoriesResponse,
    CategorizedDocumentsUploadResponse,
    DocumentGenerationRequest,
    DocumentGenerationResponse,
    DocumentGenerationStats,
    UploadedFileInfo,
    SuccessResponse,
    ErrorResponse,
    EmailRequest,
    EmailResponse,
    # Historial de documentos
    DocumentListItem,
    DocumentListResponse,
    DocumentStatsResponse,
    GetDocumentResponse,
    # Preview de documentos
    DocumentPreviewRequest,
    DocumentPreviewResponse
)
from app.services import (
    get_categories_for_type,
    DocumentGenerator,
    SessionManager,
    get_session_manager
)
from app.services.email_service import EmailService
from app.services.supabase_storage_service import SupabaseStorageService
from app.core.dependencies import get_email_service, get_supabase_storage, get_user_tenant_id, get_authenticated_user
from app.database import supabase_admin
from app.core.config import get_settings
from app.repositories.session_repository import session_repository
from app.repositories.uploaded_file_repository import uploaded_file_repository
from app.repositories.document_repository import document_repository
import hashlib

logger = structlog.get_logger()
router = APIRouter(prefix="/documents", tags=["Documents"])
settings = get_settings()

# Configuración de validación de archivos
ALLOWED_MIME_TYPES = [
    'image/jpeg',
    'image/jpg',
    'image/png',
    'image/gif',
    'image/bmp',
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # .docx
]

MAX_FILE_SIZE_BYTES = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes


# ============================================================================
# ENDPOINTS DE HISTORIAL DE DOCUMENTOS
# ============================================================================

@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    page: int = Query(1, ge=1, description="Número de página"),
    per_page: int = Query(25, ge=1, le=100, description="Documentos por página"),
    sort_by: str = Query("created_at", description="Campo para ordenar"),
    sort_order: str = Query("desc", description="Orden (asc o desc)"),
    tipo_documento: Optional[str] = Query(None, alias="type", description="Filtrar por tipo de documento"),
    estado: Optional[str] = Query(None, alias="status", description="Filtrar por estado"),
    search: Optional[str] = Query(None, description="Buscar por nombre de documento"),
    date_from: Optional[str] = Query(None, description="Fecha inicio (ISO format)"),
    date_to: Optional[str] = Query(None, description="Fecha fin (ISO format)"),
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Lista documentos generados con paginación

    Soporta filtros por tipo de documento, estado, búsqueda por nombre y rango de fechas.
    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    """
    logger.info(
        "Listando documentos",
        page=page,
        per_page=per_page,
        tenant_id=tenant_id,
        tipo_documento=tipo_documento,
        search=search,
        date_from=date_from,
        date_to=date_to
    )

    try:

        # Construir filtros simples (equality)
        filters = {}
        if tipo_documento:
            filters['tipo_documento'] = tipo_documento
        if estado:
            filters['estado'] = estado

        # Filtros avanzados (búsqueda y fechas)
        advanced_filters = {}
        if search:
            advanced_filters['search'] = search
        if date_from:
            advanced_filters['date_from'] = date_from
        if date_to:
            advanced_filters['date_to'] = date_to

        # Calcular offset
        offset = (page - 1) * per_page

        # Obtener documentos del repositorio con filtros avanzados
        documents = await document_repository.list_by_tenant_advanced(
            tenant_id=UUID(tenant_id),
            filters=filters if filters else None,
            advanced_filters=advanced_filters if advanced_filters else None,
            limit=per_page,
            offset=offset,
            order_by=sort_by,
            descending=(sort_order.lower() == 'desc')
        )

        # Obtener total para paginación
        total = await document_repository.count_by_tenant_advanced(
            tenant_id=UUID(tenant_id),
            filters=filters if filters else None,
            advanced_filters=advanced_filters if advanced_filters else None
        )

        # Calcular total de páginas
        total_pages = (total + per_page - 1) // per_page if total > 0 else 0

        # Convertir a DocumentListItem
        items = [
            DocumentListItem(
                id=str(doc['id']),
                nombre_documento=doc.get('nombre_documento', ''),
                tipo_documento=doc.get('tipo_documento', ''),
                estado=doc.get('estado', 'borrador'),
                created_at=doc['created_at'],
                storage_path=doc.get('storage_path'),
                confidence_score=doc.get('confidence_score'),
                metadata=doc.get('metadata')
            )
            for doc in documents
        ]

        logger.info(
            "Documentos listados",
            total=total,
            page=page,
            returned=len(items)
        )

        return DocumentListResponse(
            documents=items,
            total=total,
            page=page,
            per_page=per_page,
            total_pages=total_pages
        )

    except Exception as e:
        logger.error(
            "Error al listar documentos",
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al listar documentos: {str(e)}"
        )


@router.get("/stats", response_model=DocumentStatsResponse)
async def get_document_stats(
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Obtiene estadísticas de documentos del tenant

    Retorna conteos por tipo de documento y por estado.
    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    """
    logger.info("Obteniendo estadísticas de documentos", tenant_id=tenant_id)

    try:

        # Obtener todos los documentos para calcular estadísticas
        documents = await document_repository.list_by_tenant(
            tenant_id=UUID(tenant_id),
            limit=10000  # Suficiente para estadísticas
        )

        # Calcular estadísticas
        by_type = {}
        by_status = {}

        for doc in documents:
            # Contar por tipo
            tipo = doc.get('tipo_documento', 'otros')
            by_type[tipo] = by_type.get(tipo, 0) + 1

            # Contar por estado
            estado = doc.get('estado', 'borrador')
            by_status[estado] = by_status.get(estado, 0) + 1

        return DocumentStatsResponse(
            total_documents=len(documents),
            by_type=by_type,
            by_status=by_status
        )

    except Exception as e:
        logger.error(
            "Error al obtener estadísticas",
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener estadísticas: {str(e)}"
        )


# ============================================================================
# ENDPOINTS EXISTENTES
# ============================================================================

@router.get("/categories", response_model=CategoriesResponse)
async def get_categories(document_type: str):
    """
    Obtiene las 3 categorías para un tipo de documento

    Basado en por_partes.py líneas 1959-2182

    Returns:
        - parte_a: Primera categoría (Vendedor, Donador, etc.)
        - parte_b: Segunda categoría (Comprador, Donatario, etc.)
        - otros: Tercera categoría (Inmueble, Patrimonio, etc.)
    """
    logger.info("Obteniendo categorías", document_type=document_type)

    try:
        categories = get_categories_for_type(document_type)

        return CategoriesResponse(
            parte_a=categories['parte_a'],
            parte_b=categories['parte_b'],
            otros=categories['otros'],
            document_type=document_type
        )

    except Exception as e:
        logger.error(
            "Error al obtener categorías",
            document_type=document_type,
            error=str(e)
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener categorías: {str(e)}"
        )


@router.post("/upload", response_model=CategorizedDocumentsUploadResponse)
async def upload_categorized_documents(
    document_type: str = Form(...),
    template_id: str = Form(...),
    parte_a: List[UploadFile] = File(default=[]),
    parte_b: List[UploadFile] = File(default=[]),
    otros: List[UploadFile] = File(default=[]),
    session_manager: SessionManager = Depends(get_session_manager),
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Sube documentos categorizados por rol

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).

    Migración v2:
    - Usa SessionManager para almacenamiento temporal
    - Persiste en BD session_repository y uploaded_file_repository

    Args:
        document_type: Tipo de documento ('compraventa', 'donacion', etc.)
        template_id: ID del template asociado
        parte_a: Archivos de la primera categoría
        parte_b: Archivos de la segunda categoría
        otros: Archivos de la tercera categoría
        session_manager: SessionManager dependency
        tenant_id: ID del tenant (requerido)

    Returns:
        Confirmación con session_id para continuar el proceso
    """
    logger.info(
        "Upload de documentos categorizados",
        document_type=document_type,
        template_id=template_id,
        parte_a_count=len(parte_a),
        parte_b_count=len(parte_b),
        otros_count=len(otros)
    )

    try:
        # Validar que al menos haya un archivo
        total_files = len(parte_a) + len(parte_b) + len(otros)
        if total_files == 0:
            raise HTTPException(
                status_code=400,
                detail="Debe subir al menos un archivo"
            )

        # Generar ID de sesión
        session_id = f"session_{uuid.uuid4().hex[:12]}"

        # Procesar archivos por categoría
        files_info = []
        categorized_files = {
            'parte_a': [],
            'parte_b': [],
            'otros': []
        }

        async def process_category_files(files: List[UploadFile], category: str):
            for file in files:
                # VALIDACIÓN 1: Content-Type
                if file.content_type not in ALLOWED_MIME_TYPES:
                    logger.error(
                        "Tipo de archivo no permitido",
                        filename=file.filename,
                        content_type=file.content_type,
                        category=category
                    )
                    raise HTTPException(
                        status_code=400,
                        detail=f"Tipo de archivo no permitido: {file.content_type}. "
                               f"Tipos permitidos: {', '.join(ALLOWED_MIME_TYPES)}"
                    )

                # Leer contenido del archivo
                content = await file.read()

                # VALIDACIÓN 2: Tamaño máximo
                if len(content) > MAX_FILE_SIZE_BYTES:
                    logger.error(
                        "Archivo muy grande",
                        filename=file.filename,
                        size_mb=len(content) / 1024 / 1024,
                        max_size_mb=settings.MAX_FILE_SIZE_MB,
                        category=category
                    )
                    raise HTTPException(
                        status_code=413,
                        detail=f"Archivo muy grande: {len(content) / 1024 / 1024:.1f}MB. "
                               f"Tamaño máximo permitido: {settings.MAX_FILE_SIZE_MB}MB"
                    )

                # VALIDACIÓN 3: Archivo vacío
                if len(content) == 0:
                    logger.error(
                        "Archivo vacío",
                        filename=file.filename,
                        category=category
                    )
                    raise HTTPException(
                        status_code=400,
                        detail=f"El archivo '{file.filename}' está vacío"
                    )

                # Guardar info del archivo
                file_info = UploadedFileInfo(
                    filename=file.filename,
                    size_bytes=len(content),
                    content_type=file.content_type,
                    category=category
                )
                files_info.append(file_info)

                # Guardar contenido
                categorized_files[category].append({
                    'name': file.filename,
                    'content': content,
                    'content_type': file.content_type
                })

                logger.debug(
                    "Archivo procesado",
                    filename=file.filename,
                    category=category,
                    size=len(content)
                )

        # Procesar todas las categorías
        await process_category_files(parte_a, 'parte_a')
        await process_category_files(parte_b, 'parte_b')
        await process_category_files(otros, 'otros')

        # Guardar en SessionManager (almacenamiento temporal - cache rápido)
        session_manager.store_document_session(
            session_id=session_id,
            data={
                'document_type': document_type,
                'template_id': template_id,
                'categorized_files': categorized_files,
                'files_info': files_info
            }
        )

        logger.info(
            "Documentos categorizados guardados en SessionManager",
            session_id=session_id,
            total_files=total_files
        )

        # ====================================================================
        # PERSISTENCIA EN BD: Guardar sesión y archivos
        # ====================================================================
        try:
            # Crear sesión en BD (tenant_id ya validado por dependency)
            # case_id es opcional - se asigna cuando se crea/vincula un caso
            db_session = await session_repository.create_session(
                tenant_id=UUID(tenant_id),
                tipo_documento=document_type,
                total_archivos=total_files,
                session_data={
                    'template_id': template_id,
                    'session_id': session_id
                }
            )

            # Guardar metadata de archivos en BD
            for file_info in files_info:
                # Calcular hash del archivo
                file_content = next(
                    (f['content'] for f in categorized_files[file_info.category]
                     if f['name'] == file_info.filename),
                    None
                )

                if file_content:
                    file_hash = hashlib.sha256(file_content).hexdigest()

                    # TODO: Subir a Supabase Storage
                    storage_path = f"uploads/temp/{session_id}/{file_info.filename}"

                    await uploaded_file_repository.create_uploaded_file(
                        session_id=UUID(db_session['id']),
                        filename=file_info.filename,
                        storage_path=storage_path,
                        category=file_info.category,
                        original_filename=file_info.filename,
                        file_hash=file_hash,
                        content_type=file_info.content_type,
                        size_bytes=file_info.size_bytes
                    )

            logger.info(
                "Sesión y archivos persistidos en BD",
                session_id=session_id,
                db_session_id=db_session['id'],
                tenant_id=tenant_id,
                files_count=len(files_info)
            )

        except Exception as db_error:
            # No fallar el upload si falla la persistencia en BD
            logger.warning(
                "Error al persistir en BD (no crítico, datos en SessionManager)",
                error=str(db_error),
                session_id=session_id,
                tenant_id=tenant_id
            )

        return CategorizedDocumentsUploadResponse(
            session_id=session_id,
            files_received={
                'parte_a': len(parte_a),
                'parte_b': len(parte_b),
                'otros': len(otros)
            },
            total_files=total_files,
            files=files_info,
            message=f"{total_files} archivos recibidos, listos para procesar"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al procesar documentos categorizados",
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al procesar documentos: {str(e)}"
        )


@router.post("/generate", response_model=DocumentGenerationResponse)
async def generate_document(
    request: DocumentGenerationRequest,
    session_manager: SessionManager = Depends(get_session_manager),
    tenant_id: str = Depends(get_user_tenant_id),
    user: dict = Depends(get_authenticated_user),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Genera el documento final con los datos extraídos

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).

    Proceso:
    1. Obtiene el template desde SessionManager
    2. Reemplaza placeholders con los datos
    3. Aplica formato negrita (**texto**)
    4. Guarda el documento generado en SessionManager
    5. Persiste en BD y Storage
    6. Retorna URL de descarga
    """
    endpoint_start = time.time()

    logger.info(
        "generate_endpoint_starting",
        template_id=request.template_id,
        output_filename=request.output_filename,
        placeholders_count=len(request.placeholders),
        responses_count=len(request.responses),
        tenant_id=tenant_id
    )

    try:
        # Obtener template desde SessionManager (NO import circular)
        session_start = time.time()
        template_session = session_manager.get_template_session(request.template_id)
        session_duration = (time.time() - session_start) * 1000

        logger.debug("template_session_retrieved",
            template_id=request.template_id,
            found=template_session is not None,
            duration_ms=round(session_duration, 2))

        if not template_session:
            raise HTTPException(
                status_code=404,
                detail="Template no encontrado en sesión"
            )

        template_content = template_session['content']

        logger.debug("template_content_ready",
            template_size_bytes=len(template_content),
            tipo_documento=template_session.get('tipo_documento', 'unknown'))

        # === Case-insensitive matching para responses ===
        # Normaliza las claves para mapear placeholders del template con responses
        def normalize_key(key: str) -> str:
            """Normaliza una clave para matching case-insensitive"""
            return key.lower().replace(" ", "_").replace("-", "_")

        # Crear diccionario de responses ampliado con claves normalizadas
        # Esto permite que {{Placeholder}} encuentre valores de "placeholder"
        normalized_responses = dict(request.responses)  # Copia original
        for key, value in request.responses.items():
            normalized = normalize_key(key)
            # Agregar versiones normalizadas si no existen
            if normalized not in normalized_responses:
                normalized_responses[normalized] = value

        # También agregar mappings inversos: si el placeholder es "ABC" y tenemos "abc", agregar "ABC"
        for placeholder in request.placeholders:
            normalized_ph = normalize_key(placeholder)
            if placeholder not in normalized_responses and normalized_ph in normalized_responses:
                normalized_responses[placeholder] = normalized_responses[normalized_ph]

        logger.debug("responses_normalized",
            original_count=len(request.responses),
            normalized_count=len(normalized_responses))

        # Generar documento
        gen_start = time.time()
        generator = DocumentGenerator()

        document_content, stats = generator.generate_document(
            template_content=template_content,
            responses=normalized_responses,
            placeholders=request.placeholders,
            output_filename=request.output_filename
        )
        gen_duration = (time.time() - gen_start) * 1000

        logger.info("document_generated",
            doc_size_bytes=len(document_content),
            total_replaced=stats['total_replaced'],
            missing=stats['missing'],
            generation_ms=round(gen_duration, 2))

        # Generar ID para el documento
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"

        # Guardar documento en SessionManager con aislamiento por tenant
        store_start = time.time()
        session_manager.store_generated_document(
            doc_id=doc_id,
            data={
                'content': document_content,
                'filename': f"{request.output_filename}.docx",
                'size': len(document_content),
                'stats': stats
            },
            tenant_id=tenant_id
        )
        store_duration = (time.time() - store_start) * 1000

        logger.info(
            "document_stored_in_session",
            doc_id=doc_id,
            filename=request.output_filename,
            size=len(document_content),
            placeholders_replaced=stats['total_replaced'],
            store_ms=round(store_duration, 2)
        )

        # ====================================================================
        # PERSISTENCIA: Storage + BD (transaccional)
        # ====================================================================
        storage_path = None
        storage_result = None
        storage_duration = 0
        db_duration = 0

        try:
            # 1. SUBIR A SUPABASE STORAGE
            storage_start = time.time()
            logger.info("storage_upload_starting",
                doc_id=doc_id,
                tenant_id=tenant_id,
                bucket="documentos",
                size_bytes=len(document_content))

            storage_result = await supabase_storage.store_document(
                tenant_id=tenant_id,
                filename=f"{doc_id}.docx",
                content=document_content,
                metadata={
                    'doc_id': doc_id,
                    'original_filename': request.output_filename,
                    'tipo_documento': template_session.get('tipo_documento', 'desconocido')
                }
            )

            storage_duration = (time.time() - storage_start) * 1000
            storage_path = storage_result.get('path')

            logger.info(
                "storage_upload_complete",
                doc_id=doc_id,
                tenant_id=tenant_id,
                storage_path=storage_path,
                bucket=storage_result.get('bucket'),
                duration_ms=round(storage_duration, 2)
            )

            # 2. GUARDAR METADATA EN BD (solo si Storage fue exitoso)
            db_start = time.time()
            logger.debug("db_insert_starting",
                table="documentos",
                doc_id=doc_id,
                tenant_id=tenant_id)

            await document_repository.create({
                'tenant_id': tenant_id,
                'user_id': user.get('id'),  # ID del usuario autenticado
                'template_id': str(UUID(request.template_id)),
                'tipo_documento': template_session.get('tipo_documento', 'desconocido'),
                'nombre_documento': f"{request.output_filename}.docx",
                'estado': 'completado',
                'storage_path': storage_path,
                'extracted_data': request.responses,
                'metadata': {
                    'doc_id': doc_id,
                    'stats': stats,
                    'placeholders_count': len(request.placeholders),
                    'storage_bucket': storage_result.get('bucket'),
                    'storage_url': storage_result.get('url')
                }
            })

            db_duration = (time.time() - db_start) * 1000

            logger.info(
                "db_insert_complete",
                doc_id=doc_id,
                tenant_id=tenant_id,
                storage_path=storage_path,
                duration_ms=round(db_duration, 2)
            )

        except Exception as persist_error:
            # CRÍTICO: No silenciar errores de persistencia
            # Sin persistencia el documento no aparecerá en historial
            logger.error(
                "persist_failed_critical",
                error=str(persist_error),
                error_type=type(persist_error).__name__,
                doc_id=doc_id,
                tenant_id=tenant_id,
                user_id=user.get('id'),
                storage_uploaded=storage_path is not None,
                storage_ms=round(storage_duration, 2) if storage_duration else 0,
                db_ms=round(db_duration, 2) if db_duration else 0
            )
            # Re-lanzar para que el usuario sepa que hubo un problema
            raise HTTPException(
                status_code=500,
                detail=f"Documento generado pero error al guardar: {str(persist_error)}"
            )

        # Calcular lista de placeholders faltantes
        missing_list = []
        for placeholder in request.placeholders:
            value = request.responses.get(placeholder, "")
            # Un placeholder está faltante si:
            # 1. No está en responses, O
            # 2. Su valor es vacío, O
            # 3. Su valor es "No encontrado"
            if not value or value == "No encontrado" or "NO ENCONTRADO" in value.upper():
                missing_list.append(placeholder)

        # Log de placeholders faltantes si los hay
        if missing_list:
            logger.warning("placeholders_missing_detail",
                doc_id=doc_id,
                missing_count=len(missing_list),
                missing_keys=missing_list[:20])  # Limitar a 20 para no saturar logs

        # Crear response con estadísticas
        generation_stats = DocumentGenerationStats(
            placeholders_replaced=stats['total_replaced'],
            placeholders_missing=stats['missing'],
            missing_list=missing_list,
            replaced_in_body=stats.get('replaced_in_body', 0),
            replaced_in_tables=stats.get('replaced_in_tables', 0),
            replaced_in_headers=stats.get('replaced_in_headers', 0),
            replaced_in_footers=stats.get('replaced_in_footers', 0),
            bold_conversions=stats.get('bold_conversions', 0)
        )

        # Log final del endpoint con métricas totales
        total_duration = (time.time() - endpoint_start) * 1000
        logger.info("generate_endpoint_complete",
            doc_id=doc_id,
            tenant_id=tenant_id,
            filename=f"{request.output_filename}.docx",
            size_bytes=len(document_content),
            total_replaced=stats['total_replaced'],
            missing=len(missing_list),
            persisted=storage_path is not None,
            total_ms=round(total_duration, 2),
            generation_ms=round(gen_duration, 2),
            storage_ms=round(storage_duration, 2),
            db_ms=round(db_duration, 2))

        return DocumentGenerationResponse(
            success=True,
            document_id=doc_id,
            filename=f"{request.output_filename}.docx",
            download_url=f"/api/documents/download/{doc_id}",
            size_bytes=len(document_content),
            stats=generation_stats,
            message=f"Documento generado: {stats['total_replaced']} placeholders reemplazados"
        )

    except HTTPException:
        raise
    except Exception as e:
        error_duration = (time.time() - endpoint_start) * 1000
        logger.error(
            "generate_endpoint_error",
            error=str(e),
            error_type=type(e).__name__,
            template_id=request.template_id,
            tenant_id=tenant_id,
            duration_ms=round(error_duration, 2)
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al generar documento: {str(e)}"
        )


@router.post("/preview", response_model=DocumentPreviewResponse)
async def preview_document(
    request: DocumentPreviewRequest,
    session_manager: SessionManager = Depends(get_session_manager),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage),
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Genera preview del documento sin guardarlo permanentemente

    Permite al usuario ver cómo quedará el documento con los datos
    extraídos antes de confirmar la generación final.

    Proceso:
    1. Obtiene el template desde SessionManager
    2. Extrae todos los placeholders del template
    3. Reemplaza placeholders con datos proporcionados
    4. Convierte a HTML para visualización
    5. Calcula estadísticas de completitud
    6. Retorna HTML + estadísticas

    Args:
        request: DocumentPreviewRequest con template_id y data
        session_manager: SessionManager dependency

    Returns:
        DocumentPreviewResponse con HTML y estadísticas de completitud
    """
    import re
    from docx import Document
    from html import escape

    logger.info(
        "Generando preview de documento",
        template_id=request.template_id,
        data_fields=len(request.data)
    )

    try:
        # 1. Obtener template con fallback: SessionManager -> BD + Supabase Storage
        template_session = session_manager.get_template_session(request.template_id)
        template_content = None

        if template_session and template_session.get('content'):
            # Encontrado en cache de sesión
            template_content = template_session['content']
            logger.debug(
                "template_found_in_session",
                template_id=request.template_id
            )
        else:
            # Fallback: Buscar en base de datos + Supabase Storage
            logger.info(
                "template_not_in_session_trying_database",
                template_id=request.template_id,
                tenant_id=tenant_id
            )

            try:
                # Buscar en tabla templates (usar supabase_admin para bypasear RLS)
                result = supabase_admin.table('templates').select('*').eq(
                    'id', request.template_id
                ).eq('tenant_id', tenant_id).single().execute()

                if result.data:
                    template_db = result.data
                    storage_path = template_db.get('storage_path')
                    template_name = template_db.get('nombre', 'desconocido')

                    if not storage_path:
                        # Template existe en BD pero sin archivo en Storage
                        logger.warning(
                            "template_missing_storage_path",
                            template_id=request.template_id,
                            template_name=template_name,
                            created_at=template_db.get('created_at')
                        )
                        raise HTTPException(
                            status_code=404,
                            detail=f"El template '{template_name}' no tiene archivo asociado. Por favor, vuelva a subirlo desde la sección de Templates."
                        )

                    # Descargar desde Supabase Storage
                    template_content = supabase_storage.read_template(storage_path)

                    # Cachear en SessionManager para futuras peticiones
                    session_manager.store_template_session(request.template_id, {
                        'content': template_content,
                        'tipo_documento': template_db.get('tipo_documento'),
                        'filename': template_name,
                        'placeholders': template_db.get('placeholders', [])
                    })

                    logger.info(
                        "template_loaded_from_storage",
                        template_id=request.template_id,
                        storage_path=storage_path
                    )
            except HTTPException:
                # Re-raise HTTPExceptions (mensajes de error específicos)
                raise
            except Exception as db_error:
                logger.warning(
                    "template_database_lookup_failed",
                    template_id=request.template_id,
                    error=str(db_error)
                )

        if not template_content:
            raise HTTPException(
                status_code=404,
                detail="Template no encontrado en sesión ni en base de datos"
            )

        # 2. Leer documento Word para extraer texto y placeholders
        doc_stream = BytesIO(template_content)
        document = Document(doc_stream)

        # Patrón para detectar placeholders: {{nombre}} o {{nombre_campo}}
        placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')

        # Extraer todos los placeholders únicos
        all_placeholders = set()
        full_text_parts = []

        # Procesar párrafos del body
        for para in document.paragraphs:
            text = para.text
            matches = placeholder_pattern.findall(text)
            all_placeholders.update(matches)
            full_text_parts.append(text)

        # Procesar tablas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        matches = placeholder_pattern.findall(text)
                        all_placeholders.update(matches)
                        full_text_parts.append(text)

        # Procesar headers
        for section in document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text
                    matches = placeholder_pattern.findall(text)
                    all_placeholders.update(matches)
                    full_text_parts.append(text)

        # Procesar footers
        for section in document.sections:
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text
                    matches = placeholder_pattern.findall(text)
                    all_placeholders.update(matches)
                    full_text_parts.append(text)

        # 3. Calcular estadísticas
        total_placeholders = len(all_placeholders)
        filled_placeholders = 0
        missing_placeholders = []
        warnings = []

        # === Case-insensitive matching para placeholders ===
        # Normaliza las claves para evitar problemas de mayúsculas/minúsculas
        def normalize_key(key: str) -> str:
            """Normaliza una clave para matching case-insensitive"""
            return key.lower().replace(" ", "_").replace("-", "_")

        # Crear diccionario normalizado de los datos
        normalized_data = {normalize_key(k): v for k, v in request.data.items()}
        # También mantener mapeo de clave normalizada -> clave original
        key_mapping = {normalize_key(k): k for k in request.data.keys()}

        def get_value_for_placeholder(placeholder: str) -> str:
            """Busca el valor para un placeholder, usando matching case-insensitive"""
            # Primero intenta match exacto
            if placeholder in request.data:
                return request.data[placeholder]
            # Luego intenta match normalizado
            normalized = normalize_key(placeholder)
            return normalized_data.get(normalized, "")

        for placeholder in all_placeholders:
            value = get_value_for_placeholder(placeholder)
            if value and value != "No encontrado" and "NO LOCALIZADO" not in value.upper():
                filled_placeholders += 1
                # Validaciones de formato
                if placeholder.endswith('_RFC') and len(value) not in [12, 13]:
                    warnings.append(f"Campo '{placeholder}' tiene formato RFC sospechoso: {value}")
                if placeholder.endswith('_CURP') and len(value) != 18:
                    warnings.append(f"Campo '{placeholder}' tiene formato CURP sospechoso: {value}")
            else:
                missing_placeholders.append(placeholder)

        fill_percentage = (filled_placeholders / total_placeholders * 100) if total_placeholders > 0 else 0

        # 4. Generar HTML con datos reemplazados
        full_text = "\n".join(full_text_parts)

        # Reemplazar placeholders con valores (usando matching case-insensitive)
        def replace_placeholder(match):
            key = match.group(1)
            value = get_value_for_placeholder(key)
            if value and value != "No encontrado" and "NO LOCALIZADO" not in value.upper():
                return f'<span class="filled-value">{escape(value)}</span>'
            else:
                return f'<span class="missing-placeholder">{{{{{{key}}}}}}</span>'

        html_text = placeholder_pattern.sub(replace_placeholder, escape(full_text))

        # Convertir saltos de línea a <br>
        html_text = html_text.replace('\n', '<br>\n')

        # Envolver en contenedor HTML con estilos
        html_content = f'''
        <div class="document-preview">
            <style>
                .document-preview {{
                    font-family: 'Times New Roman', serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    padding: 20px;
                    background: white;
                    color: #333;
                }}
                .filled-value {{
                    background-color: #e8f5e9;
                    padding: 2px 4px;
                    border-radius: 2px;
                }}
                .missing-placeholder {{
                    background-color: #ffebee;
                    color: #c62828;
                    padding: 2px 4px;
                    border-radius: 2px;
                    font-family: monospace;
                    font-size: 10pt;
                }}
            </style>
            <div class="content">
                {html_text}
            </div>
        </div>
        '''

        logger.info(
            "Preview generado exitosamente",
            template_id=request.template_id,
            total_placeholders=total_placeholders,
            filled_placeholders=filled_placeholders,
            fill_percentage=round(fill_percentage, 1)
        )

        return DocumentPreviewResponse(
            html_content=html_content,
            total_placeholders=total_placeholders,
            filled_placeholders=filled_placeholders,
            fill_percentage=round(fill_percentage, 1),
            missing_placeholders=sorted(missing_placeholders),
            warnings=warnings
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al generar preview",
            template_id=request.template_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al generar preview: {str(e)}"
        )


@router.get("/download/{doc_id}")
async def download_document(
    doc_id: str,
    session_manager: SessionManager = Depends(get_session_manager),
    tenant_id: str = Depends(get_user_tenant_id),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Descarga un documento generado

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    Valida que el documento pertenezca al tenant antes de servir.

    Flujo:
    1. Intenta recuperar de SessionManager (cache temporal)
    2. Si no está en cache, busca en BD y descarga de Storage

    Returns:
        StreamingResponse con el archivo .docx
    """
    logger.info("Descargando documento", doc_id=doc_id, tenant_id=tenant_id)

    document_content = None
    filename = None

    # 1. Intentar recuperar de SessionManager (cache temporal)
    doc = session_manager.get_generated_document(doc_id, tenant_id=tenant_id)

    if doc:
        document_content = doc['content']
        filename = doc['filename']
        logger.info("Documento recuperado de SessionManager", doc_id=doc_id)
    else:
        # 2. Buscar en BD por doc_id (en metadata)
        try:
            db_doc = await document_repository.get_by_doc_id(doc_id, UUID(tenant_id))

            if db_doc and db_doc.get('storage_path'):
                # 3. Descargar de Supabase Storage
                storage_path = db_doc['storage_path']
                document_content = await supabase_storage.download_document(storage_path)
                filename = db_doc.get('nombre_documento', f"{doc_id}.docx")

                logger.info(
                    "Documento recuperado de Storage",
                    doc_id=doc_id,
                    storage_path=storage_path
                )
        except Exception as e:
            logger.warning(
                "Error al buscar documento en BD/Storage",
                doc_id=doc_id,
                error=str(e)
            )

    if not document_content:
        raise HTTPException(
            status_code=404,
            detail="Documento no encontrado"
        )

    # Crear stream del documento
    document_stream = BytesIO(document_content)

    logger.info(
        "Documento servido para descarga",
        doc_id=doc_id,
        filename=filename
    )

    return StreamingResponse(
        document_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


@router.post("/send-email", response_model=EmailResponse)
async def send_document_email(
    request: EmailRequest,
    email_service: EmailService = Depends(get_email_service),
    session_manager: SessionManager = Depends(get_session_manager),
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Envía un documento generado por email

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).

    Args:
        request: EmailRequest con to_email, subject, body, document_id
        email_service: EmailService dependency
        session_manager: SessionManager dependency
        tenant_id: ID del tenant (requerido)

    Returns:
        EmailResponse con confirmación del envío

    Raises:
        HTTPException 401: Si no está autenticado
        HTTPException 404: Si el documento no existe
        HTTPException 500: Si falla el envío del email
    """
    logger.info(
        "Enviando documento por email",
        to_email=request.to_email,
        document_id=request.document_id,
        tenant_id=tenant_id
    )

    try:
        # Obtener documento desde SessionManager con validación de tenant
        doc = session_manager.get_generated_document(request.document_id, tenant_id=tenant_id)

        if not doc:
            raise HTTPException(
                status_code=404,
                detail=f"Documento no encontrado en sesión: {request.document_id}"
            )

        # Preparar adjunto
        attachment_data = {
            'content': doc['content'],
            'filename': doc['filename']
        }

        # Enviar email usando EmailService (async)
        await email_service.send_email_async(
            to_email=request.to_email,
            subject=request.subject,
            body=request.body,
            attachment_data=attachment_data,
            html=request.html
        )

        logger.info(
            "Email enviado exitosamente",
            to_email=request.to_email,
            document_id=request.document_id,
            filename=doc['filename']
        )

        return EmailResponse(
            success=True,
            message=f"Email enviado exitosamente a {request.to_email}",
            to_email=request.to_email,
            document_filename=doc['filename']
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al enviar email",
            to_email=request.to_email,
            document_id=request.document_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al enviar email: {str(e)}"
        )


@router.post("/confirm/{doc_id}", response_model=SuccessResponse)
async def confirm_document(
    doc_id: str,
    tenant_id: str = Depends(get_user_tenant_id),
    session_manager: SessionManager = Depends(get_session_manager),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Confirma y guarda un documento generado en Supabase Storage

    Este endpoint debe llamarse cuando el usuario confirma que el documento
    generado es correcto y quiere guardarlo permanentemente.

    Proceso:
    1. Recupera el documento desde SessionManager
    2. Sube el documento a Supabase Storage (bucket 'documentos')
    3. Actualiza metadata en la base de datos
    4. Retorna confirmación con URL de descarga permanente

    Args:
        doc_id: ID del documento generado
        tenant_id: ID del tenant (obtenido del JWT)
        session_manager: SessionManager dependency
        supabase_storage: SupabaseStorageService dependency

    Returns:
        SuccessResponse con información del documento guardado

    Raises:
        HTTPException 401: Si no está autenticado
        HTTPException 404: Si el documento no existe en sesión
        HTTPException 500: Si falla el guardado en Supabase
    """
    logger.info(
        "Confirmando documento para guardar en Supabase",
        doc_id=doc_id,
        tenant_id=tenant_id
    )

    try:
        # Obtener documento desde SessionManager con validación de tenant
        doc = session_manager.get_generated_document(doc_id, tenant_id=tenant_id)

        if not doc:
            raise HTTPException(
                status_code=404,
                detail=f"Documento no encontrado en sesión: {doc_id}"
            )

        # Subir a Supabase Storage
        result = await supabase_storage.store_document(
            tenant_id=tenant_id,
            filename=doc['filename'],
            content=doc['content'],
            metadata={
                'doc_id': doc_id,
                'stats': doc.get('stats', {}),
                'confirmed': True
            }
        )

        logger.info(
            "Documento confirmado y guardado en Supabase",
            doc_id=doc_id,
            tenant_id=tenant_id,
            storage_path=result['path']
        )

        # Actualizar en BD con validación de tenant
        try:
            await document_repository.update_by_doc_id(
                doc_id=doc_id,
                tenant_id=UUID(tenant_id),
                updates={
                    'storage_path': result['path'],
                    'estado': 'confirmado',
                    'confirmed_at': 'now()'
                }
            )
        except Exception as db_error:
            logger.warning(
                "No se pudo actualizar en BD (no crítico)",
                error=str(db_error),
                tenant_id=tenant_id
            )

        return SuccessResponse(
            message=f"Documento confirmado y guardado exitosamente",
            data={
                'doc_id': doc_id,
                'filename': doc['filename'],
                'storage_path': result['path'],
                'url': result.get('url'),
                'bucket': result['bucket']
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al confirmar documento",
            doc_id=doc_id,
            tenant_id=tenant_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al confirmar documento: {str(e)}"
        )


# ============================================================================
# ENDPOINT PUT PARA REEMPLAZAR DOCUMENTO
# ============================================================================

@router.put("/{document_id}", response_model=SuccessResponse)
async def replace_document(
    document_id: str,
    file: UploadFile = File(..., description="Nuevo archivo .docx para reemplazar"),
    tenant_id: str = Depends(get_user_tenant_id),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Reemplaza un documento existente con uno nuevo

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    Solo permite reemplazar documentos que pertenezcan al tenant.

    Args:
        document_id: ID del documento a reemplazar (UUID)
        file: Nuevo archivo .docx
        tenant_id: ID del tenant autenticado
        supabase_storage: Servicio de Storage

    Returns:
        SuccessResponse con información del documento actualizado

    Raises:
        HTTPException 400: Si el archivo no es .docx
        HTTPException 403: Si el documento no pertenece al tenant
        HTTPException 404: Si el documento no existe
    """
    logger.info(
        "Reemplazando documento",
        document_id=document_id,
        tenant_id=tenant_id,
        new_filename=file.filename
    )

    try:
        # Validar tipo de archivo
        if not file.filename.endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail="Solo se permiten archivos .docx"
            )

        if file.content_type != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no válido: {file.content_type}"
            )

        # Obtener documento de BD
        document = await document_repository.get_by_id(UUID(document_id))

        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Documento no encontrado: {document_id}"
            )

        # Verificar que pertenece al tenant
        if str(document.get('tenant_id')) != tenant_id:
            raise HTTPException(
                status_code=403,
                detail="No tiene permisos para modificar este documento"
            )

        # Leer contenido del nuevo archivo
        new_content = await file.read()

        if len(new_content) == 0:
            raise HTTPException(
                status_code=400,
                detail="El archivo está vacío"
            )

        if len(new_content) > MAX_FILE_SIZE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Archivo muy grande. Máximo permitido: {settings.MAX_FILE_SIZE_MB}MB"
            )

        # Obtener storage_path existente o crear uno nuevo
        old_storage_path = document.get('storage_path')
        doc_id = document.get('metadata', {}).get('doc_id', document_id)

        # Subir nuevo archivo a Storage (reemplaza si existe)
        storage_result = await supabase_storage.store_document(
            tenant_id=tenant_id,
            filename=f"{doc_id}.docx",
            content=new_content,
            metadata={
                'replaced_at': datetime.now(timezone.utc).isoformat(),
                'original_filename': file.filename,
                'previous_path': old_storage_path
            }
        )

        new_storage_path = storage_result.get('path')

        # Actualizar metadata en BD
        await document_repository.update(
            document_id=UUID(document_id),
            data={
                'storage_path': new_storage_path,
                'nombre_documento': file.filename,
                'estado': 'actualizado',
                'metadata': {
                    **document.get('metadata', {}),
                    'replaced_at': datetime.now(timezone.utc).isoformat(),
                    'previous_storage_path': old_storage_path
                }
            }
        )

        logger.info(
            "Documento reemplazado exitosamente",
            document_id=document_id,
            tenant_id=tenant_id,
            new_storage_path=new_storage_path,
            old_storage_path=old_storage_path
        )

        return SuccessResponse(
            message="Documento reemplazado exitosamente",
            data={
                'document_id': document_id,
                'new_filename': file.filename,
                'storage_path': new_storage_path,
                'replaced_at': datetime.now(timezone.utc).isoformat()
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al reemplazar documento",
            document_id=document_id,
            tenant_id=tenant_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al reemplazar documento: {str(e)}"
        )


# ============================================================================
# ENDPOINTS DE ELIMINACIÓN Y EXPORTACIÓN
# ============================================================================

@router.delete("/{document_id}", response_model=SuccessResponse)
async def delete_document(
    document_id: str,
    tenant_id: str = Depends(get_user_tenant_id),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Elimina un documento del sistema

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    Solo permite eliminar documentos que pertenezcan al tenant.

    Proceso:
    1. Verifica que el documento existe y pertenece al tenant
    2. Elimina el archivo de Supabase Storage (si existe)
    3. Elimina el registro de la base de datos

    Args:
        document_id: UUID del documento a eliminar
        tenant_id: ID del tenant autenticado

    Returns:
        SuccessResponse con confirmación de eliminación
    """
    logger.info(
        "Eliminando documento",
        document_id=document_id,
        tenant_id=tenant_id
    )

    try:
        # Obtener documento para verificar propiedad y obtener storage_path
        document = await document_repository.get_by_id(UUID(document_id))

        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Documento no encontrado: {document_id}"
            )

        # Verificar que pertenece al tenant
        if str(document.get('tenant_id')) != tenant_id:
            raise HTTPException(
                status_code=403,
                detail="No tiene permisos para eliminar este documento"
            )

        # Intentar eliminar de Storage si tiene storage_path
        storage_path = document.get('storage_path')
        if storage_path:
            try:
                await supabase_storage.delete_document(storage_path)
                logger.info(
                    "Archivo eliminado de Storage",
                    storage_path=storage_path
                )
            except Exception as storage_error:
                logger.warning(
                    "Error al eliminar archivo de Storage (continuando con eliminación de BD)",
                    storage_path=storage_path,
                    error=str(storage_error)
                )

        # Eliminar de BD
        deleted = await document_repository.delete_document(
            document_id=UUID(document_id),
            tenant_id=UUID(tenant_id)
        )

        if not deleted:
            raise HTTPException(
                status_code=500,
                detail="Error al eliminar documento de la base de datos"
            )

        logger.info(
            "Documento eliminado exitosamente",
            document_id=document_id,
            tenant_id=tenant_id
        )

        return SuccessResponse(
            message="Documento eliminado exitosamente",
            data={
                'document_id': document_id,
                'deleted_at': datetime.now(timezone.utc).isoformat()
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al eliminar documento",
            document_id=document_id,
            tenant_id=tenant_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al eliminar documento: {str(e)}"
        )


@router.post("/bulk-delete", response_model=SuccessResponse)
async def bulk_delete_documents(
    document_ids: List[str] = Body(..., embed=True),
    tenant_id: str = Depends(get_user_tenant_id),
    supabase_storage: SupabaseStorageService = Depends(get_supabase_storage)
):
    """
    Elimina múltiples documentos del sistema

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    Solo elimina documentos que pertenezcan al tenant.

    Args:
        document_ids: Lista de UUIDs de documentos a eliminar
        tenant_id: ID del tenant autenticado

    Returns:
        SuccessResponse con conteo de documentos eliminados
    """
    logger.info(
        "Eliminación masiva de documentos",
        count=len(document_ids),
        tenant_id=tenant_id
    )

    if not document_ids:
        raise HTTPException(
            status_code=400,
            detail="Debe proporcionar al menos un documento para eliminar"
        )

    if len(document_ids) > 100:
        raise HTTPException(
            status_code=400,
            detail="Máximo 100 documentos por operación"
        )

    try:
        deleted_count = 0
        storage_deleted = 0
        errors = []

        for doc_id in document_ids:
            try:
                # Obtener documento para storage_path
                document = await document_repository.get_by_id(UUID(doc_id))

                if not document:
                    continue

                if str(document.get('tenant_id')) != tenant_id:
                    continue

                # Eliminar de Storage si existe
                storage_path = document.get('storage_path')
                if storage_path:
                    try:
                        await supabase_storage.delete_document(storage_path)
                        storage_deleted += 1
                    except Exception:
                        pass  # Continuar aunque falle Storage

                # Eliminar de BD
                if await document_repository.delete_document(
                    document_id=UUID(doc_id),
                    tenant_id=UUID(tenant_id)
                ):
                    deleted_count += 1

            except Exception as e:
                errors.append({'id': doc_id, 'error': str(e)})
                continue

        logger.info(
            "Eliminación masiva completada",
            requested=len(document_ids),
            deleted=deleted_count,
            storage_deleted=storage_deleted,
            errors=len(errors),
            tenant_id=tenant_id
        )

        return SuccessResponse(
            message=f"{deleted_count} de {len(document_ids)} documentos eliminados",
            data={
                'requested': len(document_ids),
                'deleted': deleted_count,
                'storage_deleted': storage_deleted,
                'errors': errors if errors else None
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error en eliminación masiva",
            tenant_id=tenant_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error en eliminación masiva: {str(e)}"
        )


@router.get("/export", response_class=StreamingResponse)
async def export_documents(
    tipo_documento: Optional[str] = Query(None, alias="type", description="Filtrar por tipo"),
    estado: Optional[str] = Query(None, alias="status", description="Filtrar por estado"),
    search: Optional[str] = Query(None, description="Buscar por nombre"),
    date_from: Optional[str] = Query(None, description="Fecha inicio"),
    date_to: Optional[str] = Query(None, description="Fecha fin"),
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Exporta documentos a CSV

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).
    Aplica los mismos filtros que el listado.

    Returns:
        StreamingResponse con archivo CSV
    """
    logger.info(
        "Exportando documentos a CSV",
        tenant_id=tenant_id,
        tipo_documento=tipo_documento,
        search=search
    )

    try:
        # Construir filtros
        filters = {}
        if tipo_documento:
            filters['tipo_documento'] = tipo_documento
        if estado:
            filters['estado'] = estado

        advanced_filters = {}
        if search:
            advanced_filters['search'] = search
        if date_from:
            advanced_filters['date_from'] = date_from
        if date_to:
            advanced_filters['date_to'] = date_to

        # Obtener todos los documentos (sin paginación)
        documents = await document_repository.list_for_export(
            tenant_id=UUID(tenant_id),
            filters=filters if filters else None,
            advanced_filters=advanced_filters if advanced_filters else None
        )

        # Generar CSV en memoria
        output = StringIO()
        writer = csv.writer(output)

        # Encabezados
        writer.writerow([
            'ID',
            'Nombre',
            'Tipo',
            'Estado',
            'Fecha Creación',
            'Fecha Actualización',
            'Score Confianza'
        ])

        # Datos
        for doc in documents:
            writer.writerow([
                doc.get('id', ''),
                doc.get('nombre_documento', ''),
                doc.get('tipo_documento', ''),
                doc.get('estado', ''),
                doc.get('created_at', ''),
                doc.get('updated_at', ''),
                doc.get('confidence_score', '')
            ])

        # Preparar respuesta
        output.seek(0)
        csv_content = output.getvalue()

        # Nombre del archivo con timestamp
        filename = f"documentos_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

        logger.info(
            "Exportación completada",
            tenant_id=tenant_id,
            documents_count=len(documents),
            filename=filename
        )

        return StreamingResponse(
            iter([csv_content]),
            media_type="text/csv",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Content-Type": "text/csv; charset=utf-8"
            }
        )

    except Exception as e:
        logger.error(
            "Error al exportar documentos",
            tenant_id=tenant_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al exportar documentos: {str(e)}"
        )


# ============================================================================
# ENDPOINT GET POR ID (debe ir al final para no capturar otras rutas)
# ============================================================================

@router.get("/{document_id}", response_model=GetDocumentResponse)
async def get_document(
    document_id: str,
    tenant_id: str = Depends(get_user_tenant_id)
):
    """
    Obtiene un documento específico por ID

    SEGURIDAD: Requiere autenticación (tenant_id obligatorio).

    Retorna toda la información del documento incluyendo datos extraídos,
    URL de descarga y metadata.

    NOTA: Este endpoint debe estar al final del router para que las rutas
    específicas como /categories, /stats, /upload tengan prioridad.
    """
    logger.info(
        "Obteniendo documento",
        document_id=document_id,
        tenant_id=tenant_id
    )

    try:

        # Obtener documento del repositorio
        document = await document_repository.get_by_id(UUID(document_id))

        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Documento no encontrado: {document_id}"
            )

        # Verificar que el documento pertenece al tenant
        if str(document.get('tenant_id')) != tenant_id:
            raise HTTPException(
                status_code=403,
                detail="No tiene permisos para acceder a este documento"
            )

        # Generar URL de descarga si hay storage_path
        download_url = None
        if document.get('storage_path'):
            download_url = f"/api/documents/download/{document_id}"

        return GetDocumentResponse(
            id=str(document['id']),
            tenant_id=str(document['tenant_id']),
            nombre_documento=document.get('nombre_documento', ''),
            tipo_documento=document.get('tipo_documento', ''),
            estado=document.get('estado', 'borrador'),
            storage_path=document.get('storage_path'),
            download_url=download_url,
            extracted_data=document.get('extracted_data'),
            edited_data=document.get('edited_data'),
            confidence_score=document.get('confidence_score'),
            metadata=document.get('metadata'),
            created_at=document['created_at'],
            updated_at=document.get('updated_at')
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Error al obtener documento",
            document_id=document_id,
            error=str(e),
            error_type=type(e).__name__
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener documento: {str(e)}"
        )
